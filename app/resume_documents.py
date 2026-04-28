from __future__ import annotations

from io import BytesIO
from pathlib import Path


SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}


def get_extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def extract_text_from_upload(filename: str, content: bytes) -> str:
    extension = get_extension(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Supported resume formats are .txt, .docx, and .pdf.")

    if extension == ".txt":
        return content.decode("utf-8", errors="replace")

    if extension == ".docx":
        from docx import Document

        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_text = []
        for table in document.tables:
            for row in table.rows:
                table_text.extend(cell.text for cell in row.cells if cell.text.strip())
        return "\n".join(paragraphs + table_text)

    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    page_text = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(page_text)
