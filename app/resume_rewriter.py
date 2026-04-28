from __future__ import annotations

import re
from io import BytesIO

from app.matcher import MatchResult


def estimate_years_experience(resume_text: str) -> int | None:
    patterns = [
        r"(\d+)\+?\s+years?\s+of\s+experience",
        r"(\d+)\+?\s+years?\s+experience",
    ]
    for pattern in patterns:
        match = re.search(pattern, resume_text.lower())
        if match:
            return int(match.group(1))
    return None


def build_rewrite_guidance(
    resume_text: str, job_description_text: str, result: MatchResult
) -> str:
    years = estimate_years_experience(resume_text)
    seniority_note = (
        f"Write for a candidate with about {years} years of experience."
        if years is not None
        else "Do not assume years of experience beyond what the resume states."
    )
    missing = ", ".join(result.missing_skills[:8]) or "No major missing skills detected."
    matched = ", ".join(result.matched_skills[:10]) or "No matched skills detected."

    return f"""Targeted Resume Rewrite Draft

Rewrite Guardrails
- Preserve truthful experience only.
- Do not invent employers, dates, tools, metrics, certifications, or responsibilities.
- {seniority_note}
- Keep language natural and candidate-written, not exaggerated or generic.
- Use the original resume as the source of truth.

Target Role Alignment
- Match score: {result.score}/100
- Readiness: {result.readiness_level}
- Matched skills: {matched}
- Skill gaps to address only if truthful: {missing}

Suggested Summary
Data engineering professional with experience building reliable data pipelines, SQL/Python workflows, and analytics-ready datasets. Interested in applying data engineering foundations to AI product workflows, including retrieval, evaluation, and production-ready data systems.

Suggested Skills Section Additions
Only add these if you have hands-on experience or if this project demonstrates them:
{chr(10).join(f"- {skill}" for skill in result.missing_skills[:8]) if result.missing_skills else "- Keep current skills and add measurable context."}

Suggested Bullet Style
- Built or improved [system/pipeline] using [tools] to support [business/user outcome].
- Automated [manual process] and improved [reliability, speed, cost, scale, or data quality].
- Designed [data workflow/API/retrieval process] with testing, monitoring, and clear ownership.

Review Checklist
- Confirm every tool listed is something you can explain in an interview.
- Add numbers only when they are true.
- Keep bullets specific to your actual projects.
- Remove any suggestion that does not match your real experience.
"""


def create_rewrite_txt(original_text: str, guidance: str) -> bytes:
    content = f"{original_text.rstrip()}\n\n\n{guidance}"
    return content.encode("utf-8")


def create_rewrite_docx(original_content: bytes, guidance: str) -> bytes:
    from docx import Document

    document = Document(BytesIO(original_content))
    document.add_page_break()
    document.add_heading("Targeted Resume Rewrite Draft", level=1)

    for line in guidance.splitlines()[2:]:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(stripped)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
